import logging
import os
import tempfile
from pathlib import Path

import pdfplumber
import pytesseract
from docx import Document
from fastapi import APIRouter, HTTPException, UploadFile
from pdf2image import convert_from_path

from backend.config import settings
from backend.db import supabase_client
from backend.models import CandidateProfile
from backend.parser.nlp_pipeline import parse_candidate

logger = logging.getLogger(__name__)

router = APIRouter()

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def _extract_pdf_text(file_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        logger.warning("pdfplumber extraction failed: %s", e)

    if not text.strip():
        logger.info("pdfplumber returned empty text, falling back to pytesseract OCR")
        text = _ocr_pdf_fallback(file_path)

    return text.strip()


def _ocr_pdf_fallback(file_path: str) -> str:
    text_parts = []
    try:
        images = convert_from_path(file_path, dpi=300)
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image, lang="ita+eng")
            if page_text.strip():
                text_parts.append(page_text)
            logger.info("OCR page %d: extracted %d chars", i + 1, len(page_text))
    except Exception as e:
        logger.error("pytesseract OCR fallback failed: %s", e)
    return "\n".join(text_parts)


def _extract_docx_text(file_path: str) -> str:
    text_parts = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
    return "\n".join(text_parts)


@router.post("/api/upload/cv")
async def upload_cv(file: UploadFile):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    ext = Path(file.filename).suffix.lower()
    if ext not in (".pdf", ".docx"):
        raise HTTPException(
            status_code=400, detail="Only PDF and DOCX files are supported"
        )

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File exceeds 10MB limit")

    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.NamedTemporaryFile(
        dir=str(upload_dir), suffix=ext, delete=False
    ) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        if ext == ".pdf":
            raw_text = _extract_pdf_text(tmp_path)
        else:
            raw_text = _extract_docx_text(tmp_path)

        if not raw_text.strip():
            raise HTTPException(
                status_code=422,
                detail="Could not extract text from the uploaded file",
            )

        profile: CandidateProfile = parse_candidate(
            raw_text=raw_text, source="cv_upload"
        )

        candidate_id = supabase_client.upsert_candidate(profile)
        profile.id = candidate_id

        return {
            "message": "CV processed successfully",
            "candidate_id": candidate_id,
            "raw_text": raw_text,
            "profile": profile.model_dump(),
        }
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
